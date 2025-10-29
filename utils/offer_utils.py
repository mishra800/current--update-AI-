from docx import Document
from pathlib import Path
from datetime import datetime
import json
from typing import Dict

OUTPUT_DIR = Path("generated_offers")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DEFAULT_TEMPLATE = """
Offer Letter

Date: {{date}}

Dear {{candidate_name}},

We are pleased to offer you the position of {{position}} at {{company_name}} starting on {{start_date}}.

Compensation: {{compensation}}

Please sign and return this letter to confirm your acceptance.

Best regards,
{{hr_name}}
{{company_name}}
"""

def generate_offer_docx(template_vars: Dict[str, str], filename_prefix: str = "offer"):
    """
    Create a simple DOCX offer letter from template_vars and return path.
    template_vars example: {candidate_name, position, start_date, compensation, hr_name, company_name}
    """
    doc = Document()
    # header paragraph
    date_str = template_vars.get("date", datetime.utcnow().strftime("%Y-%m-%d"))
    doc.add_paragraph(f"Date: {date_str}")
    doc.add_paragraph("")

    # body
    lines = [
        f"Dear {template_vars.get('candidate_name', '')},",
        "",
        f"We are pleased to offer you the position of {template_vars.get('position','')} at {template_vars.get('company_name','')} starting on {template_vars.get('start_date','')}.",
        "",
        f"Compensation: {template_vars.get('compensation','')}",
        "",
        "Please sign and return this letter to confirm your acceptance.",
        "",
        "Best regards,",
        template_vars.get("hr_name","HR Team"),
        template_vars.get("company_name","Company")
    ]
    for l in lines:
        doc.add_paragraph(l)

    stamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    fname = f"{filename_prefix}__{template_vars.get('candidate_name','candidate').replace(' ','_')}__{stamp}.docx"
    out_path = OUTPUT_DIR / fname
    doc.save(str(out_path))
    return str(out_path)
